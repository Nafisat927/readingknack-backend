from django.shortcuts import render, get_object_or_404
from django.contrib.auth import login, authenticate, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.views.decorators.csrf import ensure_csrf_cookie
from django.utils.decorators import method_decorator
from django.middleware.csrf import get_token
from django.http import JsonResponse
from passages.models import (
    UploadedDocument, QuizQuestion, QuizAnswer,
    QuizResponse, UserAnswer, GradeLevel, SkillCategory
)
from django import forms
from docx import Document
from .forms import UploadedDocumentForm
from rest_framework import viewsets, status
from rest_framework.response import Response
from rest_framework.decorators import action as drf_action
from rest_framework.views import APIView
from rest_framework.permissions import AllowAny
from .serializers import (
    UploadedDocumentSerializer, QuizQuestionSerializer, QuizAnswerSerializer,
    QuizResponseSerializer, DocumentDetailSerializer, GradeLevelSerializer, 
    SkillCategorySerializer, UserRegistrationSerializer, UserSerializer
)
import json
from .authentication import CsrfExemptSessionAuthentication
import os
from passages.gemini_utils import generate_questions, parse_questions, save_parsed_questions


# Traditional Django views (for template-based pages)
def upload_document(request):
    """Handle document upload via traditional Django form"""
    parsed_content = None

    if request.method == 'POST':
        form = UploadedDocumentForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_doc = form.save()

            doc = Document(uploaded_doc.file)
            parsed_content = "\n".join([p.text for p in doc.paragraphs])
            uploaded_doc.parsed_text = parsed_content

            try:
                print("Generating quiz questions with Gemini...")
                questions_text = generate_questions(parsed_content[:3000])
                print("Raw Gemini output:\n", questions_text)

                parsed_questions = parse_questions(questions_text)
                print("Parsed questions list:\n", parsed_questions)

                save_parsed_questions(uploaded_doc, parsed_questions)
                print("Quiz Questions saved to DB.")

            except Exception as e:
                print("Quiz generation failed:", str(e))

            uploaded_doc.save()

            return render(request, 'passages/upload_success.html', {
                'document': uploaded_doc,
                'parsed_content': parsed_content
            })
    else:
        form = UploadedDocumentForm()

    return render(request, 'passages/upload_form.html', {'form': form})


def uploaded_documents_list(request):
    """List all uploaded documents"""
    documents = UploadedDocument.objects.all()
    return render(request, 'passages/document_list.html', {'documents': documents})


def generate_questions_for_document(request, document_id):
    """Generate questions for a specific document"""
    document = get_object_or_404(UploadedDocument, id=document_id)
    parsed_text = document.parsed_text

    if not parsed_text:
        return JsonResponse({'error': 'No parsed text found in document.'}, status=400)

    questions = generate_questions(parsed_text)
    return JsonResponse({'questions': questions})


# Django REST Framework API Views
class UploadedDocumentViewSet(viewsets.ModelViewSet):
    """API endpoint for uploaded documents"""
    authentication_classes = [CsrfExemptSessionAuthentication]
    queryset = UploadedDocument.objects.all().order_by('-uploaded_at')
    serializer_class = UploadedDocumentSerializer

    def perform_create(self, serializer):
        """Handle document upload and question generation"""
        instance = serializer.save(uploader=self.request.user)
        
        # Parse .docx file
        if instance.file.name.endswith('.docx'):
            doc = Document(instance.file)
            full_text = [para.text for para in doc.paragraphs]
            instance.parsed_text = '\n'.join(full_text)
            instance.save()

        # Generate questions using AI
        try:
            parsed_text = instance.parsed_text[:3000]  # Limit to 3K chars
            print("Generating quiz questions with Gemini...")
            print("📄 Uploaded document:", instance.title)
            print("📄 Parsed text snippet:", instance.parsed_text[:300] if instance.parsed_text else "NO TEXT")

            questions_text = generate_questions(parsed_text)
            print("🤖 Raw Gemini output:\n", questions_text)

            parsed_questions = parse_questions(questions_text)
            print("🔍 Parsed questions list:", parsed_questions)

            save_parsed_questions(instance, parsed_questions)
            print("💾 Questions saved to DB.")

        except Exception as e:
            print("Quiz generation failed:", str(e))
            import traceback
            traceback.print_exc()


class DocumentDetailView(APIView):
    """Get detailed document information with questions"""
    def get(self, request, pk):
        document = get_object_or_404(UploadedDocument, pk=pk)
        serializer = DocumentDetailSerializer(document)
        return Response(serializer.data)


class QuizQuestionViewSet(viewsets.ModelViewSet):
    """API endpoint for quiz questions"""
    queryset = QuizQuestion.objects.all()
    serializer_class = QuizQuestionSerializer

    def get_queryset(self):
        document_id = self.request.query_params.get('document_id', None)
        if document_id:
            return QuizQuestion.objects.filter(document_id=document_id)
        return QuizQuestion.objects.all()


class QuizAnswerViewSet(viewsets.ModelViewSet):
    """API endpoint for quiz answers"""
    queryset = QuizAnswer.objects.all()
    serializer_class = QuizAnswerSerializer


class QuizResponseViewSet(viewsets.ModelViewSet):
    """API endpoint for quiz responses"""
    queryset = QuizResponse.objects.all()
    serializer_class = QuizResponseSerializer


class SubmitQuizView(APIView):
    """Handle quiz submission and scoring"""
    def post(self, request):
        try:
            data = request.data
            document_id = data.get('document_id')
            user_name = data.get('user_name', 'Anonymous')
            answers = data.get('answers', [])

            document = get_object_or_404(UploadedDocument, id=document_id)
            questions = QuizQuestion.objects.filter(document=document)

            if not questions.exists():
                return Response(
                    {'error': 'No questions found for this document'},
                    status=status.HTTP_400_BAD_REQUEST
                )

            # Calculate score
            score = 0
            total_questions = questions.count()
            user_answers = []

            for answer_data in answers:
                question_id = answer_data.get('question_id')
                selected_answer_id = answer_data.get('selected_answer_id')

                question = get_object_or_404(QuizQuestion, id=question_id)
                selected_answer = get_object_or_404(QuizAnswer, id=selected_answer_id)

                is_correct = selected_answer.is_correct
                if is_correct:
                    score += 1

                user_answers.append({
                    'question': question,
                    'selected_answer': selected_answer,
                    'is_correct': is_correct
                })

            # Create quiz response
            quiz_response = QuizResponse.objects.create(
                document=document,
                user_name=user_name,
                score=score,
                total_questions=total_questions
            )

            # Create user answers
            for user_answer_data in user_answers:
                UserAnswer.objects.create(
                    response=quiz_response,
                    question=user_answer_data['question'],
                    selected_answer=user_answer_data['selected_answer'],
                    is_correct=user_answer_data['is_correct']
                )

            return Response({
                'response_id': quiz_response.id,
                'score': score,
                'total_questions': total_questions,
                'percentage': round((score / total_questions) * 100, 2)
            })

        except Exception as e:
            return Response(
                {'error': str(e)},
                status=status.HTTP_400_BAD_REQUEST
            )


class GradeLevelViewSet(viewsets.ModelViewSet):
    """API endpoint for grade levels"""
    queryset = GradeLevel.objects.all()
    serializer_class = GradeLevelSerializer


class SkillCategoryViewSet(viewsets.ModelViewSet):
    """API endpoint for skill categories"""
    queryset = SkillCategory.objects.all()
    serializer_class = SkillCategorySerializer


# User Authentication API Views
class UserRegistrationView(APIView):
    """Handle user registration"""
    permission_classes = [AllowAny]
    authentication_classes = []
    
    @method_decorator(ensure_csrf_cookie)
    def get(self, request):
        """Get CSRF token for the registration form"""
        return JsonResponse({'csrfToken': get_token(request)})
    
    def post(self, request):
        """Handle user registration"""
        try:
            # Handle both form data and JSON data
            if request.content_type == 'application/json':
                data = request.data
            else:
                # Handle form data - Django UserCreationForm uses password1 and password2
                data = {
                    'username': request.POST.get('username'),
                    'password': request.POST.get('password1'),  # Django sends password1
                    'password2': request.POST.get('password2'),
                    'email': request.POST.get('email'),
                    'first_name': request.POST.get('first_name'),
                    'last_name': request.POST.get('last_name'),
                }
            
            # Validate required fields
            required_fields = ['username', 'password', 'password2', 'email', 'first_name', 'last_name']
            missing_fields = [field for field in required_fields if not data.get(field)]
            
            if missing_fields:
                return Response({
                    'success': False,
                    'error': f'Missing required fields: {", ".join(missing_fields)}'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Debug: Log the data being sent to serializer
            print(f"Registration data: {data}")
            
            serializer = UserRegistrationSerializer(data=data)
            if serializer.is_valid():
                user = serializer.save()
                
                # Log the user in after successful registration
                login(request, user)
                
                return Response({
                    'success': True,
                    'message': 'User registered successfully!',
                    'user': UserSerializer(user).data
                }, status=status.HTTP_201_CREATED)
            else:
                print(f"Serializer errors: {serializer.errors}")
                return Response({
                    'success': False,
                    'errors': serializer.errors
                }, status=status.HTTP_400_BAD_REQUEST)
                
        except Exception as e:
            print(f"Registration exception: {e}")
            return Response({
                'success': False,
                'error': str(e)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class UserLoginView(APIView):
    """Handle user login"""
    permission_classes = [AllowAny]
    authentication_classes = []
    
    @method_decorator(ensure_csrf_cookie)
    def get(self, request):
        """Get CSRF token for the login form"""
        return JsonResponse({'csrfToken': get_token(request)})
    
    def post(self, request):
        """Handle user login"""
        try:
            # Handle both form data and JSON data
            if request.content_type == 'application/json':
                username = request.data.get('username')
                password = request.data.get('password')
            else:
                # Handle form data
                username = request.POST.get('username')
                password = request.POST.get('password')
            
            if not username or not password:
                return Response({
                    'success': False,
                    'error': 'Username and password are required'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            user = authenticate(username=username, password=password)
            
            if user is not None:
                login(request, user)
                return Response({
                    'success': True,
                    'message': 'Login successful!',
                    'user': UserSerializer(user).data
                }, status=status.HTTP_200_OK)
            else:
                return Response({
                    'success': False,
                    'error': 'Invalid credentials'
                }, status=status.HTTP_401_UNAUTHORIZED)
                
        except Exception as e:
            return Response({
                'success': False,
                'error': str(e)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class UserLogoutView(APIView):
    """Handle user logout"""
    def post(self, request):
        """Handle user logout"""
        try:
            logout(request)
            return Response({
                'success': True,
                'message': 'Logout successful!'
            }, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({
                'success': False,
                'error': str(e)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class UserProfileView(APIView):
    """Get current user profile"""
    def get(self, request):
        """Get current user profile"""
        if request.user.is_authenticated:
            return Response({
                'success': True,
                'user': UserSerializer(request.user).data
            }, status=status.HTTP_200_OK)
        else:
            return Response({
                'success': False,
                'error': 'User not authenticated'
            }, status=status.HTTP_401_UNAUTHORIZED)